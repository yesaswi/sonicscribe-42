import os
import logging
from typing import Dict

import flask
from flask.typing import ResponseReturnValue
from google.cloud import storage
from pydub import AudioSegment
import functions_framework
from pydub.exceptions import CouldntDecodeError
from openai import OpenAI, OpenAIError
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class AudioTranscriberError(Exception):
    pass


class SummaryExtractorError(Exception):
    pass


class SummaryWriterError(Exception):
    pass


class OpenAIAudioTranscriber:
    def __init__(self, api_key: str, model: str):
        self.client = OpenAI(api_key=api_key)
        self.model = model

    def transcribe_audio(self, audio_file_path: str, chunk_size_ms: int = 600_000) -> str:
        try:
            with open(audio_file_path, 'rb') as audio_file:
                audio = AudioSegment.from_file(audio_file)
        except FileNotFoundError as e:
            raise AudioTranscriberError(f"Audio file not found: {audio_file_path}") from e
        except CouldntDecodeError as e:
            raise AudioTranscriberError(f"Could not decode audio file: {audio_file_path}") from e

        chunk_size = chunk_size_ms
        transcription = ""

        for i in range(0, len(audio), chunk_size):
            chunk = audio[i:i + chunk_size]
            chunk_file = "temp_chunk.wav"
            chunk.export(chunk_file, format="wav")

            with open(chunk_file, 'rb') as chunk_audio:
                try:
                    chunk_transcription = self.client.audio.transcriptions.create(
                        model=self.model,
                        file=chunk_audio,
                        response_format="json"
                    )
                    transcription += chunk_transcription.text
                except OpenAIError as e:
                    raise AudioTranscriberError(f"Error transcribing audio chunk: {e}") from e
                finally:
                    os.remove(chunk_file)

        return transcription


class OpenAISummaryExtractor:
    def __init__(self, api_key: str, model: str):
        self.client = OpenAI(api_key=api_key)
        self.model = model

    def extract_summary(self, transcription: str) -> Dict[str, str]:
        try:
            return {
                'abstract_summary': self._extract_abstract_summary(transcription),
                'key_points': self._extract_key_points(transcription),
                'sentiment': self._analyze_sentiment(transcription)
            }
        except OpenAIError as e:
            raise SummaryExtractorError(f"Error extracting summary: {e}") from e

    def _extract_abstract_summary(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are a highly skilled AI trained in language comprehension and summarization. I would "
                           "like you to read the following text and summarize it into a concise abstract paragraph. "
                           "Aim to retain the most important points, providing a coherent and readable summary that "
                           "could help a person understand the main points of the discussion without needing to read "
                           "the entire text. Please avoid unnecessary details or tangential points.",
            user_content=transcription
        )

    def _extract_key_points(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are a proficient AI with a specialty in distilling information into key points. Based "
                           "on the following text, identify and list the main points that were discussed or brought "
                           "up. These should be the most important ideas, findings, or topics that are crucial to the "
                           "essence of the discussion. Your goal is to provide a list that someone could read to "
                           "quickly understand what was talked about.",
            user_content=transcription
        )

    def _analyze_sentiment(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="As an AI with expertise in language and emotion analysis, your task is to analyze the "
                           "sentiment of the following text. Please consider the overall tone of the discussion, "
                           "the emotion conveyed by the language used, and the context in which words and phrases are "
                           "used. Indicate whether the sentiment is generally positive, negative, or neutral, "
                           "and provide brief explanations for your analysis where possible.",
            user_content=transcription
        )

    def _chat_completion(self, system_content: str, user_content: str) -> str:
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                temperature=0,
                messages=[
                    {"role": "system", "content": system_content},
                    {"role": "user", "content": user_content}
                ]
            )
            return response.choices[0].message.content
        except OpenAIError as e:
            raise SummaryExtractorError(f"Error in chat completion: {e}") from e


class MeetingMinutesExtractor(OpenAISummaryExtractor):
    def extract_summary(self, transcription: str) -> Dict[str, str]:
        summary = super().extract_summary(transcription)
        summary['action_items'] = self._extract_action_items(transcription)
        return summary

    def _extract_action_items(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are an AI expert in analyzing conversations and extracting action items. Please "
                           "review the text and identify any tasks, assignments, or actions that were agreed upon or "
                           "mentioned as needing to be done. These could be tasks assigned to specific individuals, "
                           "or general actions that the group has decided to take. Please list these action items "
                           "clearly and concisely.",
            user_content=transcription
        )


class PodcastSummaryExtractor(OpenAISummaryExtractor):
    def extract_summary(self, transcription: str) -> Dict[str, str]:
        summary = super().extract_summary(transcription)
        summary['topics_covered'] = self._extract_topics_covered(transcription)
        summary['guest_info'] = self._extract_guest_info(transcription)
        return summary

    def _extract_topics_covered(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="As an AI trained in content analysis, your task is to review the following podcast "
                           "transcription and list the main topics that were covered in the episode. These should be "
                           "the key themes, subjects, or ideas that were discussed at length. Please provide a clear "
                           "and concise list of these topics.",
            user_content=transcription
        )

    def _extract_guest_info(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are an AI assistant skilled in identifying and summarizing information about guests "
                           "mentioned in a podcast. Based on the provided transcript, please list the names and "
                           "relevant information (if available) about any guests that were part of the podcast episode."
                           "This could include their background, expertise, or the main points they contributed to the "
                           "discussion.",
            user_content=transcription
        )


class YouTubeSummaryExtractor(OpenAISummaryExtractor):
    def extract_summary(self, transcription: str) -> Dict[str, str]:
        summary = super().extract_summary(transcription)
        summary['main_topics'] = self._extract_main_topics(transcription)
        summary['takeaways'] = self._extract_takeaways(transcription)
        return summary

    def _extract_main_topics(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="As an AI content analyst, your job is to review the following YouTube video transcription "
                           "and identify the main topics covered in the video. These should be the key subjects, "
                           "themes, or ideas that the video focuses on and discusses in detail. Please provide a "
                           "clear and concise list of these main topics.",
            user_content=transcription
        )

    def _extract_takeaways(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are an AI assistant with expertise in summarizing key learnings from video content. "
                           "Based on the provided YouTube video transcript, please list the main takeaways or lessons "
                           "that a viewer could learn from watching this video. These should be the most important "
                           "points, insights, or conclusions that the video aims to convey.",
            user_content=transcription
        )


class SummaryWriter:
    @staticmethod
    def save_as_docx(summary: Dict[str, str], filename: str) -> None:
        try:
            doc = Document()
            for key, value in summary.items():
                heading = ' '.join(word.capitalize() for word in key.split('_'))
                doc.add_heading(heading, level=1)
                doc.add_paragraph(value)
                doc.add_paragraph()
            doc.save(filename)
        except Exception as e:
            raise SummaryWriterError(f"Error saving summary: {e}") from e


@functions_framework.http
def process_audio(request: flask.Request) -> ResponseReturnValue:
    # For more information about CORS and CORS preflight requests, see:
    # https://developer.mozilla.org/en-US/docs/Glossary/Preflight_request

    # Set CORS headers for the preflight request
    if request.method == "OPTIONS":
        # Allows GET requests from any origin with the Content-Type
        # header and caches preflight response for a 3600s
        headers = {
            "Access-Control-Allow-Origin": "https://sonicscribe-m3bfm4czka-uk.a.run.app",
            "Access-Control-Allow-Methods": "GET",
            "Access-Control-Allow-Headers": "Content-Type",
            "Access-Control-Max-Age": "3600",
        }

        return "", 204, headers

    # Set CORS headers for the main request
    headers = {"Access-Control-Allow-Origin": "https://sonicscribe-m3bfm4czka-uk.a.run.app"}

    try:
        # Extract necessary parameters from the request
        audio_file = request.files.get('audio')
        if audio_file is None:
            return {'error': 'Audio file not provided'}, 400

        audio_type = request.form.get('type', 'meeting').lower()
        if audio_type not in ['meeting', 'youtube', 'podcast']:
            return {'error': 'Invalid audio type. Please choose from: meeting, youtube, or podcast.'}, 400

        model = request.form.get('model', 'gpt-4-turbo')
        if model not in ['gpt-4-turbo', 'gpt-4o', 'gpt-3.5-turbo']:
            return {'error': 'Invalid model. Please choose from: gpt-4-turbo, gpt-4o, gpt-3.5-turbo.'}, 400

        access_code = request.form.get('access_code')
        if access_code is None:
            return {'error': 'Access code not provided'}, 400
        if access_code != os.environ.get('ACCESS_CODE'):
            return {'error': 'Invalid access code'}, 403

        # Construct the filename for the audio file in Cloud Storage
        audio_filename = audio_file.filename

        # Save the audio file to a temporary location
        temp_audio_path = f'/tmp/{audio_filename}'
        audio_file.save(temp_audio_path)

        # Validate the audio file format
        try:
            audio = AudioSegment.from_file(temp_audio_path)
        except CouldntDecodeError as e:
            logger.error(f"Could not decode audio file: {temp_audio_path} - {e}")
            return {'error': 'Invalid audio file format'}, 400

        # Save the audio file to Cloud Storage
        storage_client = storage.Client()
        bucket_name = os.environ.get('AUDIO_BUCKET_NAME')
        if bucket_name is None:
            logger.error("STORAGE_BUCKET environment variable not set")
            return {'error': 'Internal Server Error'}, 500

        bucket = storage_client.bucket(bucket_name)
        audio_blob = bucket.blob(audio_filename)
        audio_blob.upload_from_filename(temp_audio_path)

        # Transcribe the audio
        openai_api_key = os.environ.get('OPENAI_API_KEY')
        if openai_api_key is None:
            logger.error("OPENAI_API_KEY environment variable not set")
            return {'error': 'Internal Server Error'}, 500

        transcriber = OpenAIAudioTranscriber(openai_api_key, "whisper-1")
        try:
            transcription = transcriber.transcribe_audio(temp_audio_path)
        except AudioTranscriberError as e:
            logger.error(str(e))
            return {'error': 'Error during audio transcription'}, 500

        # Extract the summary
        if audio_type == "meeting":
            summary_extractor = MeetingMinutesExtractor(openai_api_key, model)
        elif audio_type == "youtube":
            summary_extractor = YouTubeSummaryExtractor(openai_api_key, model)
        elif audio_type == "podcast":
            summary_extractor = PodcastSummaryExtractor(openai_api_key, model)
        else:
            return {'error': 'Invalid audio type. Please choose from: meeting, youtube, or podcast.'}, 400

        try:
            summary = summary_extractor.extract_summary(transcription)
        except SummaryExtractorError as e:
            logger.error(str(e))
            return {'error': 'Error during summary extraction'}, 500

        # Save the summary as a Word document
        output_filename = f"{os.path.splitext(audio_filename)[0]}_summary.docx"
        output_file_path = f"/tmp/{output_filename}"
        try:
            SummaryWriter.save_as_docx(summary, output_file_path)
        except SummaryWriterError as e:
            logger.error(str(e))
            return {'error': 'Error during summary generation'}, 500

        # Upload the summary file to Cloud Storage
        summary_blob = bucket.blob(output_filename)
        summary_blob.upload_from_filename(output_file_path)

        # Remove the temporary summary file
        os.remove(output_file_path)

        # Remove the temporary audio file
        os.remove(temp_audio_path)

        return {'transcription': transcription, 'summary': summary}, 200, headers

    except Exception as e:
        logger.exception(f'An error occurred during audio processing - {e}')
        return {'error': 'Internal Server Error'}, 500
